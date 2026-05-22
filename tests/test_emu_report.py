"""Tests for the EMU admission packet builder.

Stage (a) covers: header (patient block + allergy callout + failed ASMs),
current regimen + rescue protocol, data-quality footer.
"""

from __future__ import annotations

import zipfile
from datetime import date
from pathlib import Path

import pytest
from docx import Document

from bennett_care.emu_report import EmuInputs, build_emu_doc
from bennett_care.ingest import load_log


@pytest.fixture
def emu_inputs(sample_log_path: Path) -> EmuInputs:
    log = load_log(sample_log_path)
    return EmuInputs(
        log=log,
        admission_date=date(2026, 7, 15),
        lookback_days=30,
    )


def _all_text(doc) -> str:
    """Concatenate all paragraph and table cell text in document order."""
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_build_writes_valid_docx(emu_inputs: EmuInputs, tmp_path: Path) -> None:
    out = tmp_path / "emu.docx"
    written = build_emu_doc(emu_inputs, output_path=out)
    assert written == out
    assert written.exists()
    assert zipfile.is_zipfile(written)


def test_brand_band_text(emu_inputs: EmuInputs, tmp_path: Path) -> None:
    out = tmp_path / "emu.docx"
    build_emu_doc(emu_inputs, output_path=out)
    text = _all_text(Document(out))
    assert "Bennett" in text
    assert "EMU Admission Packet" in text
    assert "Admission: 2026-07-15" in text


def test_patient_facts_block(emu_inputs: EmuInputs, tmp_path: Path) -> None:
    out = tmp_path / "emu.docx"
    build_emu_doc(emu_inputs, output_path=out)
    text = _all_text(Document(out))
    assert "Patient" in text
    assert "35 lb" in text
    assert "15.9 kg" in text
    assert "EMAS" in text
    assert "Lennox-Gastaut" in text
    assert "Children's Healthcare of Atlanta" in text


def test_allergy_callout_present(emu_inputs: EmuInputs, tmp_path: Path) -> None:
    out = tmp_path / "emu.docx"
    build_emu_doc(emu_inputs, output_path=out)
    text = _all_text(Document(out))
    assert "ALLERGY" in text
    assert "Valproate" in text
    assert "skin reaction" in text
    assert "April 2026" in text
    assert "Do not order" in text


def test_failed_asms_line(emu_inputs: EmuInputs, tmp_path: Path) -> None:
    out = tmp_path / "emu.docx"
    build_emu_doc(emu_inputs, output_path=out)
    text = _all_text(Document(out))
    assert "Failed prior ASMs" in text
    for drug in ("prednisolone", "topiramate", "felbamate", "valproate"):
        assert drug in text


def test_current_regimen_appears(emu_inputs: EmuInputs, tmp_path: Path) -> None:
    """The latest real med change in the fixture is 2026-01-05 (Epidiolex)."""
    out = tmp_path / "emu.docx"
    build_emu_doc(emu_inputs, output_path=out)
    text = _all_text(Document(out))
    assert "Current regimen" in text
    assert "2026-01-05" in text
    assert "Epidiolex" in text


def test_rescue_protocol_listed(emu_inputs: EmuInputs, tmp_path: Path) -> None:
    out = tmp_path / "emu.docx"
    build_emu_doc(emu_inputs, output_path=out)
    text = _all_text(Document(out))
    assert "Rescue protocol" in text
    assert "Valtoco" in text
    assert "Klonopin" in text


def test_data_quality_footer(emu_inputs: EmuInputs, tmp_path: Path) -> None:
    out = tmp_path / "emu.docx"
    build_emu_doc(emu_inputs, output_path=out)
    text = _all_text(Document(out))
    assert "Data-quality footer" in text
    assert "Monitored days" in text
    assert "excluded" in text.lower()
    assert "mismatch" in text.lower()
    assert "Seizure Type coverage" in text


def test_no_causal_language_in_footer(emu_inputs: EmuInputs, tmp_path: Path) -> None:
    """The footer explicitly disclaims causation; sanity-check phrasing."""
    out = tmp_path / "emu.docx"
    build_emu_doc(emu_inputs, output_path=out)
    text = _all_text(Document(out))
    assert "No causal claims" in text
