"""Tests for bennett_care.report."""

from __future__ import annotations

import zipfile
from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
import pytest
from docx import Document

from bennett_care.ingest import Dose, MedChange, SeizureLog
from bennett_care.report import ReportInputs, build_report
from bennett_care.stats import analyze_recent_changes
from bennett_care.visualize import render_all_charts


def _make_log(days: int = 200) -> SeizureLog:
    rng = np.random.default_rng(0)
    dates = pd.date_range("2026-01-01", periods=days, freq="D")
    counts = rng.poisson(lam=20, size=days)
    counts[150:] = rng.poisson(lam=8, size=days - 150)  # post-change drop
    daily = pd.DataFrame(
        {
            "count": counts,
            "daily_total_recorded": counts,
            "mismatch": np.zeros(days, dtype=bool),
        },
        index=dates,
    )
    daily.index.name = "Date"

    cluster_count = 60
    cluster_dates = [dates[i] for i in range(0, days, days // cluster_count)][:cluster_count]
    clusters = pd.DataFrame(
        {
            "Date": cluster_dates,
            "Cluster #": list(range(1, cluster_count + 1)),
            "Start Time": ["8:00 AM"] * cluster_count,
            "Duration (min)": [1.0] * cluster_count,
            "Seizure Count": [1] * cluster_count,
            "Seizure Type": (["atonic"] * 30 + [""] * 30)[:cluster_count],
            "Day Status": ["logged"] * cluster_count,
            "Verified": ["Y"] * cluster_count,
            "Flags": ["rescue_meds_given"] * 3 + [None] * (cluster_count - 3),
            "Notes": [None] * cluster_count,
            "flag_tokens": (
                [frozenset({"rescue_meds_given"})] * 3
                + [frozenset()] * (cluster_count - 3)
            ),
            "hour": [8.0] * cluster_count,
        }
    )

    changes = [
        MedChange(date=dates[0], raw="Clobazam 1mL am",
                  regimen={"Clobazam": (Dose(1.0, "am"),)}),
        MedChange(date=dates[100], raw="Clobazam 2mL am",
                  regimen={"Clobazam": (Dose(2.0, "am"),)}),
        MedChange(date=dates[150], raw="Clobazam 3mL am",
                  regimen={"Clobazam": (Dose(3.0, "am"),)}),
    ]
    return SeizureLog(
        daily=daily,
        clusters=clusters,
        med_changes=changes,
        mismatches=pd.DataFrame(),
        excluded_dates=pd.DatetimeIndex([]),
    )


@pytest.fixture
def built_report(tmp_path: Path) -> Path:
    log = _make_log()
    charts = render_all_charts(log, output_dir=tmp_path / "charts")
    analyses = analyze_recent_changes(log, k=2)
    inputs = ReportInputs(
        log=log,
        chart_paths=charts,
        analyses=analyses,
        visit_date=date(2026, 8, 4),
        lookback_days=90,
    )
    out = tmp_path / "report.docx"
    build_report(inputs, output_path=out)
    return out


def test_report_is_valid_docx_zip(built_report: Path) -> None:
    assert built_report.exists()
    assert built_report.stat().st_size > 5000
    assert zipfile.is_zipfile(built_report)
    with zipfile.ZipFile(built_report) as zf:
        assert "[Content_Types].xml" in zf.namelist()
        assert any(name.startswith("word/media/") for name in zf.namelist())


def test_report_contains_brand_band_and_visit_date(built_report: Path) -> None:
    doc = Document(str(built_report))
    text = "\n".join(p.text for t in doc.tables for row in t.rows for cell in row.cells for p in cell.paragraphs)
    text += "\n" + "\n".join(p.text for p in doc.paragraphs)
    assert "Bennett" in text
    assert "Pre-Visit Summary" in text
    assert "2026-08-04" in text  # visit date


def test_report_section_headings_present(built_report: Path) -> None:
    doc = Document(str(built_report))
    body_text = "\n".join(p.text for p in doc.paragraphs)
    for heading in [
        "1. Patient & current regimen",
        "2. Weekly seizure totals & medication-change timeline",
        "3. 14-day rolling average",
        "4. Seizures by hour of day",
        "5. Pre / post comparison",
        "6. Seizure type distribution",
        "7. Flag summary",
        "8. Open clinical questions",
        "9. Appendix",
    ]:
        assert heading in body_text, f"Missing heading: {heading}"


def test_report_does_not_have_notable_days_section(built_report: Path) -> None:
    doc = Document(str(built_report))
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Notable days" not in body_text


def test_report_med_change_table_present(built_report: Path) -> None:
    """Section 2's horizontal med-change timeline table."""
    doc = Document(str(built_report))
    med_tables = [
        t for t in doc.tables
        if t.rows[0].cells[0].text.strip() == "Date"
        and t.rows[0].cells[1].text.strip() == "Change"
        and "regimen" in t.rows[0].cells[2].text.lower()
    ]
    assert len(med_tables) == 1


def test_report_reading_key_present(built_report: Path) -> None:
    """Section 5's plain-English glossary."""
    doc = Document(str(built_report))
    body_text = "\n".join(p.text for p in doc.paragraphs)
    # Glossary intro line
    assert "How to read this table" in body_text
    # Magnitude labels in Hedges' g should appear in at least one cell.
    table_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert any(label in table_text for label in ("small", "medium", "large", "very large", "trivial"))


def test_report_pvalue_has_plain_english_suffix(built_report: Path) -> None:
    doc = Document(str(built_report))
    table_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    # _fmt_pvalue always appends "noise"-keyworded suffix when p is not None
    assert "noise" in table_text.lower()


def test_report_current_regimen_renders(built_report: Path) -> None:
    doc = Document(str(built_report))
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Current regimen" in body_text
    assert "Clobazam" in body_text


def test_report_prepost_table_has_three_window_rows_per_change(built_report: Path) -> None:
    doc = Document(str(built_report))
    # Find tables whose first row starts with "Window".
    window_tables = [
        t for t in doc.tables if t.rows[0].cells[0].text.strip() == "Window"
    ]
    assert len(window_tables) == 2  # two recent changes
    for t in window_tables:
        # 1 header + 3 windows (14/28/56)
        assert len(t.rows) == 4


def test_report_appendix_table_has_30_data_rows(built_report: Path) -> None:
    doc = Document(str(built_report))
    appendix_tables = [
        t for t in doc.tables
        if t.rows[0].cells[0].text.strip() == "Date"
        and "Day" in t.rows[0].cells[1].text
    ]
    assert len(appendix_tables) == 1
    assert len(appendix_tables[0].rows) == 1 + 30  # header + 30 days


def test_report_flag_table_present(built_report: Path) -> None:
    doc = Document(str(built_report))
    flag_tables = [
        t for t in doc.tables
        if t.rows[0].cells[0].text.strip() == "Event type"
    ]
    assert len(flag_tables) == 1
    # Header + 4 event-type rows (school events explicitly removed per user request).
    assert len(flag_tables[0].rows) == 5


def test_report_flag_table_omits_school_events(built_report: Path) -> None:
    doc = Document(str(built_report))
    flag_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
        if t.rows[0].cells[0].text.strip() == "Event type"
    )
    assert "school" not in flag_text.lower()


def test_report_section_1_shows_lookback_window_not_monitored_days(built_report: Path) -> None:
    """Section 1's facts table should show the analysis window, NOT 'Monitored days' or 'Excluded dates'."""
    doc = Document(str(built_report))
    # The first table after the brand band is section 1's facts table; find it by its left-column labels.
    facts_table = None
    for t in doc.tables:
        labels = [t.rows[i].cells[0].text.strip() for i in range(len(t.rows))]
        if "Patient" in labels and "Visit date" in labels:
            facts_table = t
            break
    assert facts_table is not None
    labels = {t_row.cells[0].text.strip() for t_row in facts_table.rows}
    assert "Charts cover" in labels
    assert "Monitored days in log" not in labels
    assert "Excluded dates (unmonitored / uncertain)" not in labels
    # The "Charts cover" value should mention the day count and a date range.
    value_for_charts = next(
        row.cells[1].text.strip() for row in facts_table.rows
        if row.cells[0].text.strip() == "Charts cover"
    )
    assert "90 days" in value_for_charts
    assert "–" in value_for_charts or "-" in value_for_charts


def test_report_no_causal_language(built_report: Path) -> None:
    """Document must not contain phrases that ASSERT causation per CLAUDE.md.

    Disclaimer language (e.g. "not what caused it") is fine and important — we ban
    affirmative causal claims, not the word "caused" in any context.
    """
    doc = Document(str(built_report))
    all_text = " ".join(p.text for p in doc.paragraphs).lower()
    all_text += " " + " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    ).lower()
    banned_assertions = [
        "the drug worked",
        "responded well",
        "reduced seizures",
        "the change caused",
        "the drug caused",
        "due to the change",
        "effective treatment",
    ]
    for phrase in banned_assertions:
        assert phrase not in all_text, f"Forbidden causal phrase in report: {phrase!r}"
