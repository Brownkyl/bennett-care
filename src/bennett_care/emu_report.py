"""Assemble the EMU admission packet Word document.

Different audience from the outpatient visit-prep doc: this is for the EMU floor
team and the admitting neurologist on duty. The clinical question is *"what
should we expect during capture, and what are the safe orders right now?"* — so
the document is distributional, not temporal.

Section layout:
  Brand band: "Bennett — EMU Admission Packet | Admission: <date>"
  1. Patient & admission summary (incl. Valproate allergy + failed prior ASMs)
  2. Current regimen + rescue protocol
  3. Recent seizure profile — distributional (durations, sizes, types)
  4. Seizures by hour of day (reuse — critical for capture timing)
  5. Rescue events table (last N days)
  6. Medication change history (last 6 months)
  7. Data-quality footer

All output is factual; no causal language. No temporal trend chart by design
(see ``bennett-care-charts-as-implicit-claims`` decision in memory).
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell

from .ingest import SeizureLog
from .report import (
    BODY_FONT,
    BODY_PT,
    BRAND_COLOR_HEX,
    HEADING_PT,
    INSTITUTION,
    NEUROLOGIST,
    PATIENT_NAME,
    PATIENT_WEIGHT_KG,
    PATIENT_WEIGHT_LB,
    _add_caption,
    _add_section_heading,
    _configure_default_style,
    _configure_page,
    _set_cell_background,
    _set_cell_padding,
    _set_cell_text,
)

# --- EMU-specific patient constants (from CLAUDE.md "Patient context") --- #
DIAGNOSIS: str = (
    "EMAS (epilepsy with myoclonic-atonic seizures), "
    "possibly evolving toward Lennox-Gastaut"
)
VALPROATE_ALLERGY_DATE: str = "April 2026"
VALPROATE_ALLERGY_NOTE: str = "skin reaction; formally documented"
FAILED_PRIOR_ASMS: tuple[str, ...] = (
    "prednisolone",
    "topiramate",
    "felbamate",
    "valproate",
)
RESCUE_PROTOCOL: tuple[tuple[str, str], ...] = (
    ("Valtoco", "intranasal diazepam"),
    ("Klonopin ODT", "oral clonazepam"),
)

ALLERGY_BANNER_HEX: str = "C0504D"  # muted red for the allergy callout


@dataclass(frozen=True)
class EmuInputs:
    """Inputs for ``build_emu_doc``.

    Stage (a) carries only the data the header + footer need; chart paths and
    derived analyses are added by subsequent stages.
    """

    log: SeizureLog
    admission_date: date
    lookback_days: int = 180


def build_emu_doc(inputs: EmuInputs, *, output_path: str | Path) -> Path:
    """Build the EMU admission packet .docx and write it to ``output_path``."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure_page(doc)
    _configure_default_style(doc)

    _add_brand_band(doc, inputs.admission_date)
    _add_section_1_patient(doc, inputs)
    _add_allergy_callout(doc)
    _add_failed_asms_line(doc)
    _add_section_2_regimen_rescue(doc, inputs)
    _add_data_quality_footer(doc, inputs)

    doc.save(output_path)
    return output_path


# --------------------------------------------------------------------------- #
# Brand band                                                                  #
# --------------------------------------------------------------------------- #


def _add_brand_band(doc: DocxDocument, admission_date: date) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_background(cell, BRAND_COLOR_HEX)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        f"{PATIENT_NAME} — EMU Admission Packet     |     "
        f"Admission: {admission_date.isoformat()}"
    )
    run.font.name = BODY_FONT
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_padding(cell, top=120, bottom=120)
    doc.add_paragraph()


# --------------------------------------------------------------------------- #
# Section 1: Patient & admission summary                                      #
# --------------------------------------------------------------------------- #


def _add_section_1_patient(doc: DocxDocument, inputs: EmuInputs) -> None:
    _add_section_heading(doc, "1. Patient & admission")
    facts = [
        ("Patient", f"{PATIENT_NAME} (~{PATIENT_WEIGHT_LB:.0f} lb / {PATIENT_WEIGHT_KG:.1f} kg)"),
        ("Diagnosis", DIAGNOSIS),
        ("Admission date", inputs.admission_date.isoformat()),
        ("Data through", inputs.log.latest_date.date().isoformat()),
        ("Outpatient neurologist", f"{NEUROLOGIST}, {INSTITUTION}"),
    ]
    table = doc.add_table(rows=len(facts), cols=2)
    table.autofit = False
    for i, (label, value) in enumerate(facts):
        row = table.rows[i]
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.3)
        _set_cell_text(row.cells[0], label, bold=True)
        _set_cell_text(row.cells[1], value)


def _add_allergy_callout(doc: DocxDocument) -> None:
    """Highlighted allergy banner. Must be visually unmissable for EMU intake."""
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_background(cell, ALLERGY_BANNER_HEX)
    _set_cell_padding(cell, top=100, bottom=100)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    label = para.add_run("ALLERGY  ")
    label.font.bold = True
    label.font.size = Pt(12)
    label.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    detail = para.add_run(
        f"Valproate — {VALPROATE_ALLERGY_NOTE} ({VALPROATE_ALLERGY_DATE}). "
        "Do not order valproate or VPA-containing regimens."
    )
    detail.font.size = Pt(11)
    detail.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _add_failed_asms_line(doc: DocxDocument) -> None:
    doc.add_paragraph()
    para = doc.add_paragraph()
    head = para.add_run("Failed prior ASMs: ")
    head.bold = True
    para.add_run(", ".join(FAILED_PRIOR_ASMS) + ".")


# --------------------------------------------------------------------------- #
# Section 2: Current regimen + rescue protocol                                #
# --------------------------------------------------------------------------- #


def _add_section_2_regimen_rescue(doc: DocxDocument, inputs: EmuInputs) -> None:
    _add_section_heading(doc, "2. Current regimen & rescue protocol")
    real_changes = [c for c in inputs.log.med_changes if c.regimen]

    sub = doc.add_paragraph()
    if real_changes:
        latest = real_changes[-1]
        run = sub.add_run(
            f"Current regimen (as of {latest.date.date().isoformat()}):"
        )
        run.bold = True
        for drug, doses in latest.regimen.items():
            dose_str = ", ".join(f"{d.amount_ml} mL {d.timing}" for d in doses)
            doc.add_paragraph(f"  • {drug}: {dose_str}")
    else:
        run = sub.add_run("No parsed regimen found in the log.")
        run.italic = True

    doc.add_paragraph()
    rescue_para = doc.add_paragraph()
    rescue_head = rescue_para.add_run("Rescue protocol:")
    rescue_head.bold = True
    for name, desc in RESCUE_PROTOCOL:
        doc.add_paragraph(f"  • {name} ({desc})")


# --------------------------------------------------------------------------- #
# Data-quality footer                                                         #
# --------------------------------------------------------------------------- #


def _add_data_quality_footer(doc: DocxDocument, inputs: EmuInputs) -> None:
    doc.add_paragraph()
    _add_section_heading(doc, "Data-quality footer")
    log = inputs.log
    end = log.latest_date
    start = end - pd.Timedelta(days=inputs.lookback_days - 1)
    in_window = log.daily.loc[(log.daily.index >= start) & (log.daily.index <= end)]
    excluded_in_window = log.excluded_dates[
        (log.excluded_dates >= start) & (log.excluded_dates <= end)
    ]
    mismatches_in_window = log.mismatches.loc[
        (log.mismatches.index >= start) & (log.mismatches.index <= end)
    ]
    typed_clusters = int((log.clusters["Seizure Type"].astype(str).str.strip() != "").sum())
    total_clusters = int(len(log.clusters))
    typed_pct = (100 * typed_clusters / total_clusters) if total_clusters else 0.0

    para = doc.add_paragraph()
    head = para.add_run(
        f"Window: last {inputs.lookback_days} days "
        f"({start.strftime('%b %-d, %Y')} – {end.strftime('%b %-d, %Y')})."
    )
    head.bold = True

    bullets = [
        f"Monitored days in window: {len(in_window)}.",
        f"Days excluded (Day Status unmonitored or uncertain): {len(excluded_in_window)}.",
        f"Daily-Total / cluster-sum mismatches in window: {len(mismatches_in_window)} — "
        "cluster-sum is the source of truth.",
        f"Seizure Type coverage: {typed_clusters} of {total_clusters} clusters typed "
        f"({typed_pct:.1f}%). Type distribution reflects this typed slice only.",
    ]
    for b in bullets:
        doc.add_paragraph(f"  • {b}")

    _add_caption(
        doc,
        "Counts are descriptive. No causal claims about regimen or events are made in "
        "this document.",
    )
