"""Assemble the Phase 1 pre-visit summary Word document.

Section layout (per CLAUDE.md):
  Brand band: "Bennett — Pre-Visit Summary | Visit: <date>"
  1. Header (patient, visit, neurologist, current regimen)
  2. Weekly seizure totals chart + horizontal medication-change timeline table
  3. 14-day rolling average chart
  4. Seizures by hour of day (24-bar chart, no day-of-week dimension)
  5. Pre/post statistics with a "How to read this" key block (plain-English
     column headers, magnitude labels on Hedges' g, explicit no-causation note)
  6. Seizure type distribution
  7. Flags summary (rescue, ended-in-tonic, school)
  8. Open clinical questions (empty template)
  9. Appendix: raw daily counts for last 30 days

All output is factual; no causal language. The user composes narrative
from this document during the clinic visit.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell

from .ingest import SeizureLog, count_flag
from .stats import (
    PrePostAnalysis,
    format_regimen,
    format_regimen_delta,
    real_changes_with_prev,
    rescue_event_dates,
)
from .visualize import ChartPaths

# --- Patient constants (Phase 1; later we may externalize to a config file) ---
PATIENT_NAME: str = "Bennett"
PATIENT_WEIGHT_LB: float = 35.0
PATIENT_WEIGHT_KG: float = 15.9
NEUROLOGIST: str = "Dr. Anna Lecticia Ribeiro-Pinto"
INSTITUTION: str = "Children's Healthcare of Atlanta"

# --- Styling ---
BRAND_COLOR_HEX: str = "1F4E79"  # dark blue
BODY_FONT: str = "Calibri"
BODY_PT: int = 11
HEADING_PT: int = 13
APPENDIX_DAYS: int = 30


@dataclass(frozen=True)
class ReportInputs:
    """Everything build_report needs. Bundling keeps the call site readable."""

    log: SeizureLog
    chart_paths: ChartPaths
    analyses: list[PrePostAnalysis]
    visit_date: date
    lookback_days: int = 90


def build_report(inputs: ReportInputs, *, output_path: str | Path) -> Path:
    """Build the pre-visit summary .docx and write it to ``output_path``."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure_page(doc)
    _configure_default_style(doc)

    _add_brand_band(doc, inputs.visit_date)
    _add_section_1_header(doc, inputs)
    _add_section_2_trends(doc, inputs)
    _add_chart_section(doc, "3. 14-day rolling average", inputs.chart_paths.rolling_avg,
                       caption="14-day trailing mean of daily seizure counts. Smooths week-to-week "
                               "noise so longer-term trajectories are visible.")
    _add_chart_section(doc, "4. Seizures by hour of day", inputs.chart_paths.hour_distribution,
                       caption=f"Total seizures by hour across the last {inputs.lookback_days} days "
                               "of monitored clusters with a recorded time. Peak hour is highlighted.")
    _add_section_5_prepost(doc, inputs)
    _add_chart_section(doc, "6. Seizure type distribution", inputs.chart_paths.type_distribution,
                       caption="Distribution among clusters that have a Seizure Type recorded "
                               "(coverage shown in chart).")
    _add_section_7_flags(doc, inputs)
    _add_section_8_questions(doc)
    _add_section_9_appendix(doc, inputs)

    doc.save(output_path)
    return output_path


# --------------------------------------------------------------------------- #
# Page / style                                                                #
# --------------------------------------------------------------------------- #


def _configure_page(doc: DocxDocument) -> None:
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.8)


def _configure_default_style(doc: DocxDocument) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)


# --------------------------------------------------------------------------- #
# Brand band                                                                  #
# --------------------------------------------------------------------------- #


def _add_brand_band(doc: DocxDocument, visit_date: date) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_background(cell, BRAND_COLOR_HEX)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"{PATIENT_NAME} — Pre-Visit Summary     |     Visit: {visit_date.isoformat()}")
    run.font.name = BODY_FONT
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Vertical padding on the band
    _set_cell_padding(cell, top=120, bottom=120)
    # Spacer paragraph below the band
    doc.add_paragraph()


def _set_cell_background(cell: _Cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    tc_pr.append(shading)


def _set_cell_padding(cell: _Cell, *, top: int = 80, bottom: int = 80) -> None:
    """Padding values are in twentieths of a point (dxa). 120 ≈ 6pt."""
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("bottom", bottom)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tc_pr.append(margins)


# --------------------------------------------------------------------------- #
# Section helpers                                                             #
# --------------------------------------------------------------------------- #


def _add_section_heading(doc: DocxDocument, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(HEADING_PT)
    run.font.bold = True
    # Bottom border on the paragraph
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_caption(doc: DocxDocument, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def _add_chart_section(doc: DocxDocument, heading: str, image_path: Path, *, caption: str) -> None:
    _add_section_heading(doc, heading)
    doc.add_picture(str(image_path), width=Inches(6.5))
    _add_caption(doc, caption)


# --------------------------------------------------------------------------- #
# Section 1: Header                                                           #
# --------------------------------------------------------------------------- #


def _add_section_1_header(doc: DocxDocument, inputs: ReportInputs) -> None:
    _add_section_heading(doc, "1. Patient & current regimen")
    log = inputs.log

    facts = [
        ("Patient", f"{PATIENT_NAME} (~{PATIENT_WEIGHT_LB:.0f} lb / {PATIENT_WEIGHT_KG:.1f} kg)"),
        ("Visit date", inputs.visit_date.isoformat()),
        ("Data through", log.latest_date.date().isoformat()),
        ("Monitored days in log", f"{len(log.daily)}"),
        ("Excluded dates (unmonitored / uncertain)", f"{len(log.excluded_dates)}"),
        ("Neurologist", f"{NEUROLOGIST}, {INSTITUTION}"),
    ]
    table = doc.add_table(rows=len(facts), cols=2)
    table.autofit = False
    for i, (label, value) in enumerate(facts):
        row = table.rows[i]
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.3)
        _set_cell_text(row.cells[0], label, bold=True)
        _set_cell_text(row.cells[1], value)

    # Current regimen
    doc.add_paragraph()
    real_changes = [c for c in log.med_changes if c.regimen]
    if real_changes:
        latest = real_changes[-1]
        sub = doc.add_paragraph()
        run = sub.add_run(f"Current regimen (as of {latest.date.date().isoformat()}):")
        run.bold = True
        for drug, doses in latest.regimen.items():
            dose_str = ", ".join(f"{d.amount_ml} mL {d.timing}" for d in doses)
            doc.add_paragraph(f"  • {drug}: {dose_str}")
    else:
        doc.add_paragraph("No parsed regimen found in the log.")


def _set_cell_text(cell: _Cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_PT)
    if bold:
        run.bold = True


# --------------------------------------------------------------------------- #
# Section 2: Trends + medication-change timeline                              #
# --------------------------------------------------------------------------- #


def _add_section_2_trends(doc: DocxDocument, inputs: ReportInputs) -> None:
    _add_section_heading(doc, "2. Weekly seizure totals & medication-change timeline")
    doc.add_picture(str(inputs.chart_paths.weekly_trend), width=Inches(6.5))
    _add_caption(
        doc,
        f"Weekly sums of daily seizure counts over the last {inputs.lookback_days} days. "
        "Dotted line marks the median weekly total in this window.",
    )
    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Medication changes in this window:")
    sub_run.bold = True

    log = inputs.log
    end = log.latest_date
    pd_mod = __import__("pandas")
    start = end - pd_mod.Timedelta(days=inputs.lookback_days - 1)
    pairs = real_changes_with_prev(log)
    in_window = [(c, p) for c, p in pairs if start <= c.date <= end]

    if not in_window:
        doc.add_paragraph(f"  No regimen changes in the last {inputs.lookback_days} days.")
        return

    headers = ["Date", "Change", "Full regimen after change"]
    table = doc.add_table(rows=1 + len(in_window), cols=3)
    table.autofit = False
    widths = (Inches(1.0), Inches(2.4), Inches(3.1))
    for j, (h, w) in enumerate(zip(headers, widths)):
        _set_cell_text(table.rows[0].cells[j], h, bold=True)
        table.rows[0].cells[j].width = w
    for i, (change, prev) in enumerate(in_window):
        cells = table.rows[i + 1].cells
        delta = format_regimen_delta(prev.regimen if prev else None, change.regimen)
        regimen_str = format_regimen(change.regimen)
        _set_cell_text(cells[0], change.date.strftime("%b %-d, %Y"))
        _set_cell_text(cells[1], delta)
        _set_cell_text(cells[2], regimen_str)
        for cell, w in zip(cells, widths):
            cell.width = w


# --------------------------------------------------------------------------- #
# Section 5: Pre/post statistics with a plain-English key block               #
# --------------------------------------------------------------------------- #


def _add_section_5_prepost(doc: DocxDocument, inputs: ReportInputs) -> None:
    _add_section_heading(doc, "5. Pre / post comparison — two most recent regimen changes")
    if not inputs.analyses:
        doc.add_paragraph("No real regimen changes in the log; nothing to analyze.")
        return

    intro = doc.add_paragraph(
        "This section compares Bennett's daily seizure counts in the days BEFORE and "
        "AFTER each of his two most recent regimen changes. Three window sizes are shown "
        "on either side of each change (14, 28, and 56 days) so short-term and longer-term "
        "patterns are both visible. All numbers are descriptive — they say what happened in "
        "time around the change, not what caused it."
    )
    intro.paragraph_format.space_after = Pt(6)

    _add_reading_key(doc)
    doc.add_paragraph()

    for analysis in inputs.analyses:
        title = doc.add_paragraph()
        t_run = title.add_run(
            f"Change on {analysis.change_date.strftime('%B %-d, %Y')} — {analysis.delta_str}"
        )
        t_run.bold = True
        t_run.font.size = Pt(12)

        sub = doc.add_paragraph()
        sub_run = sub.add_run(f"Full regimen after this change: {analysis.regimen_str}")
        sub_run.italic = True
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        headers = [
            "Window",
            "Avg daily seizures BEFORE (95% CI)",
            "Days observed",
            "Avg daily seizures AFTER (95% CI)",
            "Days observed",
            "Effect size (Hedges' g)",
            "How likely is this from chance? (p)",
            "Another change inside window?",
        ]
        table = doc.add_table(rows=1 + len(analysis.windows), cols=len(headers))
        for j, h in enumerate(headers):
            _set_cell_text(table.rows[0].cells[j], h, bold=True)
        for i, days in enumerate(sorted(analysis.windows)):
            w = analysis.windows[days]
            row = table.rows[i + 1].cells
            _set_cell_text(row[0], f"{days} days")
            _set_cell_text(row[1], _fmt_mean_ci(w.pre))
            _set_cell_text(row[2], str(w.pre.n))
            _set_cell_text(row[3], _fmt_mean_ci(w.post))
            _set_cell_text(row[4], str(w.post.n))
            _set_cell_text(row[5], _fmt_hedges(w.hedges_g))
            _set_cell_text(row[6], _fmt_pvalue(w.mann_whitney_p))
            _set_cell_text(row[7], _fmt_overlap(w.pre_contains_other_change,
                                                w.post_contains_other_change))
        doc.add_paragraph()

    caveat = doc.add_paragraph()
    c_run = caveat.add_run(
        "Important: these comparisons are observational. Other factors that may also "
        "have changed (sleep, illness, growth, other medications, season, the two recent "
        "changes contaminating each other's windows) could contribute to any difference. "
        "The statistics describe what happened, not why."
    )
    c_run.italic = True
    c_run.font.size = Pt(9.5)
    c_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def _add_reading_key(doc: DocxDocument) -> None:
    """A small 2-column glossary so the table cells are readable without prior stats knowledge."""
    para = doc.add_paragraph()
    head = para.add_run("How to read this table:")
    head.bold = True

    rows = [
        (
            "Avg daily seizures (95% CI)",
            "Average number of seizures per day in the window, with the 95% "
            "confidence interval — the range of plausible true averages given the data. "
            "Computed by bootstrap (10,000 resamples).",
        ),
        (
            "Days observed",
            "Number of monitored days actually present in the window. May be fewer "
            "than the window length when data runs out or some days were excluded.",
        ),
        (
            "Effect size (Hedges' g)",
            "How big the difference between BEFORE and AFTER is, measured in units of "
            "typical day-to-day variation. Conventional labels: 0.2 = small, 0.5 = medium, "
            "0.8 = large, 1.2 or more = very large. A negative number means AFTER is LOWER "
            "than BEFORE.",
        ),
        (
            "How likely is this from chance? (p)",
            "Mann-Whitney U two-sided p-value. The probability of seeing a difference at "
            "least this big purely by random fluctuation if BEFORE and AFTER actually had "
            "the same underlying rate. Smaller = less likely to be noise alone. No "
            "significance threshold is applied — this is descriptive, not a hypothesis test.",
        ),
        (
            "Another change inside window?",
            "\"pre\", \"post\", or both: another regimen change falls inside this window, "
            "so the comparison is contaminated and the effect attributed to THIS change "
            "may partly reflect the other one. Treat with caution.",
        ),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, (term, definition) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].width = Inches(2.0)
        cells[1].width = Inches(4.5)
        _set_cell_text(cells[0], term, bold=True)
        _set_cell_text(cells[1], definition)


def _fmt_mean_ci(ci) -> str:
    if ci.n == 0:
        return "—"
    if ci.low == ci.high == ci.mean:
        return f"{ci.mean:.2f}"
    return f"{ci.mean:.2f} [{ci.low:.2f}, {ci.high:.2f}]"


def _fmt_hedges(g: float | None) -> str:
    if g is None:
        return "n/a"
    mag = abs(g)
    if mag < 0.2:
        label = "trivial"
    elif mag < 0.5:
        label = "small"
    elif mag < 0.8:
        label = "medium"
    elif mag < 1.2:
        label = "large"
    else:
        label = "very large"
    return f"{g:+.2f} ({label})"


def _fmt_pvalue(p: float | None) -> str:
    if p is None:
        return "n/a"
    if p < 0.001:
        plain = "very unlikely to be noise alone"
    elif p < 0.01:
        plain = "unlikely to be noise alone"
    elif p < 0.05:
        plain = "somewhat unlikely to be noise"
    elif p < 0.10:
        plain = "could be noise"
    else:
        plain = "consistent with noise"
    return f"{p:.4f} — {plain}"


def _fmt_overlap(pre: bool, post: bool) -> str:
    parts = []
    if pre:
        parts.append("pre")
    if post:
        parts.append("post")
    return ", ".join(parts) if parts else "no"


# --------------------------------------------------------------------------- #
# Section 7: Flags                                                            #
# --------------------------------------------------------------------------- #


def _add_section_7_flags(doc: DocxDocument, inputs: ReportInputs) -> None:
    _add_section_heading(doc, "7. Flag summary")
    log = inputs.log
    end = log.latest_date
    start = end - __import__("pandas").Timedelta(days=inputs.lookback_days - 1)
    in_window = log.clusters[(log.clusters["Date"] >= start) & (log.clusters["Date"] <= end)]

    rescue_clusters = count_flag(in_window, "rescue_meds_given")
    rescue_notes = sum(
        1
        for c in log.med_changes
        if not c.regimen
        and "rescue" in c.raw.lower()
        and start <= c.date <= end
    )
    rescue_dates = rescue_event_dates(log)
    rescue_dates_in_window = sum(1 for d in rescue_dates if start <= d <= end)
    ended_tonic = count_flag(in_window, "ended_in_tonic")
    school = count_flag(in_window, ["school_data", "school_data_pending"])

    headers = ["Event type", "Count in last 90 days", "Notes"]
    rows = [
        ("Rescue meds given (cluster flags)", rescue_clusters, "Clusters tagged `rescue_meds_given`."),
        ("Rescue notes in Meds column", rescue_notes, "Free-text rescue entries co-located with regimen log."),
        ("Distinct dates with rescue evidence", rescue_dates_in_window, "Union of the two rows above."),
        ("Ended in tonic", ended_tonic, "Clusters tagged `ended_in_tonic`."),
        ("School events", school, "Clusters tagged `school_data` or `school_data_pending`."),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    for j, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[j], h, bold=True)
    for i, (label, count, note) in enumerate(rows):
        cells = table.rows[i + 1].cells
        _set_cell_text(cells[0], label)
        _set_cell_text(cells[1], str(count))
        _set_cell_text(cells[2], note)


# --------------------------------------------------------------------------- #
# Section 8: Open questions                                                   #
# --------------------------------------------------------------------------- #


def _add_section_8_questions(doc: DocxDocument) -> None:
    _add_section_heading(doc, "8. Open clinical questions")
    para = doc.add_paragraph()
    r = para.add_run("(Add questions here before the visit.)")
    r.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    for _ in range(5):
        doc.add_paragraph("• ")


# --------------------------------------------------------------------------- #
# Section 9: Appendix                                                         #
# --------------------------------------------------------------------------- #


def _add_section_9_appendix(doc: DocxDocument, inputs: ReportInputs) -> None:
    _add_section_heading(doc, f"9. Appendix — raw daily counts, last {APPENDIX_DAYS} days")
    pd = __import__("pandas")
    daily = inputs.log.daily
    end = inputs.log.latest_date
    start = end - pd.Timedelta(days=APPENDIX_DAYS - 1)
    window = daily.loc[(daily.index >= start) & (daily.index <= end)]

    headers = ["Date", "Day", "Count", "Recorded Daily Total", "Mismatch?"]
    table = doc.add_table(rows=1 + len(window), cols=len(headers))
    for j, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[j], h, bold=True)
    for i, (idx, row) in enumerate(window.iterrows()):
        cells = table.rows[i + 1].cells
        _set_cell_text(cells[0], idx.date().isoformat())
        _set_cell_text(cells[1], idx.strftime("%a"))
        _set_cell_text(cells[2], str(int(row["count"])))
        _set_cell_text(cells[3], str(int(row["daily_total_recorded"])))
        _set_cell_text(cells[4], "yes" if row["mismatch"] else "")
